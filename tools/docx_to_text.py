"""
docx_to_text.py — Converteer een .docx bestand naar Markdown-tekst.

Gebruik:
    python3 tools/docx_to_text.py --input pad/naar/bestand.docx
    python3 tools/docx_to_text.py --input pad/naar/bestand.docx --output .tmp/origineel.txt

Output:
    Tekst wordt geprint naar stdout én opgeslagen in .tmp/origineel.txt (standaard).
    Inline opmaak wordt bewaard: **vet** en *cursief* als Markdown-markers.
    Dit is relevant voor de APA-check: cursieve titels in de literatuurlijst blijven zichtbaar.

Afbeeldingen:
    Inline afbeeldingen worden geëxtraheerd naar <output_dir>/images/ als PNG-bestanden.
    In de tekst wordt een Markdown-figuurplaceholder geplaatst: ![caption](pad/naar/afbeelding.png)
    Dit maakt round-trip herschrijven mogelijk: md_to_docx.py plaatst de afbeeldingen terug.
"""

import argparse
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def format_runs(para) -> str:
    """
    Combineer alle runs van een alinea met inline Markdown-opmaak.
    Vet → **tekst**, cursief → *tekst*.
    Aangrenzende runs met dezelfde opmaak worden samengevoegd.
    """
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        if is_bold and is_italic:
            parts.append(f'***{text}***')
        elif is_bold:
            parts.append(f'**{text}**')
        elif is_italic:
            parts.append(f'*{text}*')
        else:
            parts.append(text)
    return ''.join(parts).strip()


_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
_NS_V = 'urn:schemas-microsoft-com:vml'


def _tag(ns: str, name: str) -> str:
    return f'{{{ns}}}{name}'


def para_has_image(para) -> bool:
    """Controleer of een alinea een inline afbeelding bevat (w:drawing of v:imagedata)."""
    xml = para._element
    return bool(
        xml.findall('.//' + _tag(_NS_W, 'drawing')) or
        xml.findall('.//' + _tag(_NS_V, 'imagedata'))
    )


def extract_image_rids(para):
    """Haal alle relationship-IDs op van inline afbeeldingen in een alinea."""
    xml = para._element
    rids = []
    # Inline drawing (modern OOXML) — a:blip r:embed
    for blip in xml.findall('.//' + _tag(_NS_A, 'blip')):
        rid = blip.get(_tag(_NS_R, 'embed'))
        if rid:
            rids.append(rid)
    # Legacy VML imagedata r:id
    for img in xml.findall('.//' + _tag(_NS_V, 'imagedata')):
        rid = img.get(_tag(_NS_R, 'id'))
        if rid:
            rids.append(rid)
    return rids


def _table_to_markdown(table) -> list:
    """Converteer een python-docx Table naar markdown-tabelregels."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return []
    md_lines = []
    # Eerste rij als header
    header = rows[0]
    md_lines.append('| ' + ' | '.join(header) + ' |')
    md_lines.append('| ' + ' | '.join('---' for _ in header) + ' |')
    for row in rows[1:]:
        # Zorg dat rij evenveel kolommen heeft als header
        padded = row + [''] * (len(header) - len(row))
        md_lines.append('| ' + ' | '.join(padded[:len(header)]) + ' |')
    return md_lines


def extract_text(docx_path: str, images_dir: Path = None) -> str:
    doc = Document(docx_path)
    lines = []
    image_counter = 0

    # Bouw lookup-maps van XML-element → python-docx object (behoudt correcte parent-chain)
    para_map = {para._element: para for para in doc.paragraphs}
    table_map = {table._element: table for table in doc.tables}

    # Itereer over body-elementen in documentvolgorde (paragrafen én tabellen)
    body = doc.element.body
    for child in body:
        # --- Tabel ---
        if child.tag == qn('w:tbl'):
            table = table_map.get(child)
            if table:
                md_lines = _table_to_markdown(table)
                if md_lines:
                    lines.append('')
                    lines.extend(md_lines)
                    lines.append('')
            continue

        # --- Paragraaf ---
        if child.tag != qn('w:p'):
            continue

        para = para_map.get(child)
        if not para:
            continue
        style = para.style.name if para.style else ''

        # Detecteer alinea met afbeelding
        if para_has_image(para):
            rids = extract_image_rids(para)
            saved_paths = []

            if images_dir is not None:
                for rid in rids:
                    try:
                        image_part = doc.part.related_parts.get(rid)
                        if image_part is None:
                            continue
                        image_counter += 1
                        # Bepaal extensie op basis van content-type
                        content_type = getattr(image_part, 'content_type', '')
                        if 'png' in content_type:
                            ext = '.png'
                        elif 'jpeg' in content_type or 'jpg' in content_type:
                            ext = '.jpg'
                        elif 'gif' in content_type:
                            ext = '.gif'
                        elif 'svg' in content_type:
                            ext = '.svg'
                        else:
                            ext = '.png'  # fallback
                        img_filename = f'figure_{image_counter:02d}{ext}'
                        img_path = images_dir / img_filename
                        img_path.write_bytes(image_part.blob)
                        saved_paths.append(str(img_path))
                    except Exception:
                        pass

            # Haal eventuele tekst uit dezelfde alinea op (alt-text of caption-run)
            inline_text = format_runs(para).strip()

            if saved_paths:
                # Gebruik eerste opgeslagen afbeelding; caption volgt in volgende alinea
                caption_text = inline_text or ''
                lines.append(f'![{caption_text}]({saved_paths[0]})')
            elif rids:
                # Afbeelding aanwezig maar niet opgeslagen (geen images_dir)
                lines.append(f'![{inline_text}](afbeelding_{image_counter + 1})')
                image_counter += 1
            continue

        text = format_runs(para)

        if not text:
            lines.append('')
            continue

        if 'Heading 1' in style:
            lines.append(f'# {text}')
        elif 'Heading 2' in style:
            lines.append(f'## {text}')
        elif 'Heading 3' in style:
            lines.append(f'### {text}')
        elif 'List' in style or para.style.name.startswith('List'):
            lines.append(f'- {text}')
        else:
            lines.append(text)

    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='Converteer .docx naar Markdown-tekst')
    parser.add_argument('--input', required=True, help='Pad naar het .docx bestand')
    parser.add_argument(
        '--output', default='.tmp/origineel.txt',
        help='Uitvoerbestand (standaard: .tmp/origineel.txt)'
    )
    parser.add_argument(
        '--no-images', action='store_true',
        help='Afbeeldingen NIET extraheren (alleen tekst)'
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'Fout: bestand niet gevonden: {input_path}', file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != '.docx':
        print(f'Fout: verwacht een .docx bestand, niet {input_path.suffix}', file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output)
    output_path.parent.mkdir(exist_ok=True)

    # Afbeeldingenmap naast het outputbestand
    if args.no_images:
        images_dir = None
    else:
        images_dir = output_path.parent / 'images'
        images_dir.mkdir(exist_ok=True)

    text = extract_text(str(input_path), images_dir=images_dir)

    output_path.write_text(text, encoding='utf-8')

    print(text)
    print(f'\n[Opgeslagen in {output_path}]', file=sys.stderr)
    if images_dir and any(images_dir.iterdir()):
        n = sum(1 for _ in images_dir.iterdir())
        print(f'[{n} afbeelding(en) opgeslagen in {images_dir}/]', file=sys.stderr)


if __name__ == '__main__':
    main()
