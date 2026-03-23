"""
word_export.py — APA 7 .docx generator voor DutchQuill AI

Gebruik:
    python3 tools/word_export.py --input .tmp/export_payload.json --output .tmp/rapport.docx

Stdout bij succes: absoluut pad naar het gegenereerde .docx-bestand
Stderr + exit code 1 bij fout: foutmelding

===========================================================================
JSON-SCHEMA (volledig overzicht)
===========================================================================

Verplichte sleutels zijn gemarkeerd met [VERPLICHT]. Alle andere zijn optioneel.

{
  "metadata": {                        [VERPLICHT]
    "title": "Rapporttitel",           [VERPLICHT] — ook de kop van de inleiding
    "subtitle": "Ondertitel",
    "authors": ["Naam Achternaam"],
    "student_numbers": ["123456"],
    "institution": "Naam Instelling",
    "faculty": "Faculteitnaam",
    "course": "Vaknaam — Vakcode",
    "supervisor": "Dr. Naam",
    "submission_date": "19 maart 2026",
    "font": "times"                    — "times" | "arial" | "georgia" (standaard: "times")
  },

  "abstract": {                        — Aparte pagina na titelblad
    "text": "Samenvattingstekst...",
    "keywords": ["trefwoord1", "trefwoord2"]
  },

  "abbreviations": [                   — Lijst van afkortingen; weggelaten = geen pagina
    {"afkorting": "MKB", "definitie": "Midden- en Kleinbedrijf"}
  ],

  "introduction_text": [ <blokken> ],  — Inleiding; kop = metadata.title
  "body":               [ <blokken> ],  — Alle hoofdstuktekst inclusief koppen
  "conclusion_text":    [ <blokken> ],  — Conclusie; kop = "Conclusie"

  "references": [                      — APA-referenties als strings; *cursief* via *...*;
    "Auteur, A. (2023). *Titel.* Uitgever."  DOI-URLs worden automatisch hyperlink
  ],

  "appendices": [                      — Na literatuurlijst; elk op nieuwe pagina
    {
      "label": "Bijlage A",            — Kop niveau 1 (verschijnt in inhoudsopgave)
      "title": "Beschrijvende titel",  — Gecentreerd vet, normal-stijl (niet in TOC)
      "content": [ <blokken> ]
    }
  ]
}

---------------------------------------------------------------------------
BLOKTYPEN (voor introduction_text, body, conclusion_text en appendix content)
---------------------------------------------------------------------------

Alinea:
  {"type": "paragraph", "text": "Tekst van de alinea."}
  — Altijd eerste-regel-inspringing (1,27 cm), dubbele regelafstand

Kop:
  {"type": "heading", "level": 1, "text": "Hoofdstuktitel"}
  — level 1: gecentreerd, vet; nieuwe pagina in body
  — level 2: links, vet
  — level 3: links, vet cursief
  — level 4: ingesprongen, vet, punt + inline tekst op zelfde regel
    {"type": "heading", "level": 4, "text": "Kopnaam", "inline_text": "Tekst na de kop."}
  — level 5: ingesprongen, vet cursief, punt + inline tekst
    {"type": "heading", "level": 5, "text": "Kopnaam", "inline_text": "Tekst na de kop."}

Blokcitaat (≥40 woorden, APA):
  {"type": "block_quote", "text": "Citaattekst...", "citation": "(Auteur, jaar, p. X)"}

Tabel:
  {
    "type": "table",
    "number": 1,
    "title": "Beschrijvende tabeltitel",
    "headers": ["Kolom A", "Kolom B"],
    "rows": [["Waarde 1", "Waarde 2"]],
    "note": "Noot onder de tabel."     — optioneel; wordt "Noot. ..." met cursief label
  }
  — APA-opmaak: nummer + titel bóven, geen verticale lijnen, noot ónder

Figuurplaatshouder:
  {
    "type": "figure_placeholder",
    "number": 1,
    "caption": "Figuuronderschrift.",
    "image_path": ".tmp/figuur_1.jpg",  — optioneel; .jpg, .jpeg of .png; bestand moet bestaan
    "note": "Noot onder het figuur."    — optioneel
  }
  — Met image_path: beeldinbedding via add_picture(); zonder: grijze plaatshouder
  — APA 7: label + titel staan altijd bóven de afbeelding; noot staat ónder

Pagina-einde:
  {"type": "page_break"}

Codeblok (terminal-output, broncode, commandoregel):
  {"type": "code", "tekst": "nmap -sV 192.168.1.1\nPORT   STATE SERVICE"}
  — Courier New 11pt, grijze achtergrond (#D9D9D9), links inspringen 1 cm, enkelvoudige regelafstand
  — Elke regel = eigen alinea, zodat achtergrond doorloopt als blok
  — Niet APA-standaard, maar goedgekeurd voor gebruik in DutchQuill AI rapporten

---------------------------------------------------------------------------
COMPACT VOORBEELDPAYLOAD (minimaal werkend document)
---------------------------------------------------------------------------

{
  "metadata": {
    "institution": "Naam Instelling",
    "title": "Titel van het Rapport",
    "authors": ["Voornaam Achternaam"],
    "student_numbers": ["123456"],
    "course": "Vaknaam — VAK301",
    "supervisor": "Dr. M. van den Berg",
    "submission_date": "19 maart 2026",
    "font": "times"
  },
  "abstract": {
    "text": "Samenvattingstekst van 150–250 woorden.",
    "keywords": ["trefwoord1", "trefwoord2"]
  },
  "introduction_text": [
    {"type": "paragraph", "text": "Eerste alinea van de inleiding."}
  ],
  "body": [
    {"type": "heading", "level": 1, "text": "Theoretisch Kader"},
    {"type": "paragraph", "text": "Eerste alinea van het theoretisch kader."}
  ],
  "conclusion_text": [
    {"type": "paragraph", "text": "Conclusie van het onderzoek."}
  ],
  "references": [
    "Auteur, A. A. (2023). *Titel van het werk.* Uitgever."
  ],
  "appendices": []
}

Volledig werkend voorbeeld: zie .tmp/test_payload.json
===========================================================================
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Waarschuwingsverzameling (gevuld tijdens documentopbouw, geprint naar stdout)
# ---------------------------------------------------------------------------

_build_warnings: list = []


# ---------------------------------------------------------------------------
# Lettertypeconfiguratie
# ---------------------------------------------------------------------------

FONT_CONFIG = {
    "times": {"name": "Times New Roman", "size": Pt(12)},
    "arial": {"name": "Arial", "size": Pt(11)},
    "georgia": {"name": "Georgia", "size": Pt(11)},
}


def get_font_config(font_key: str) -> dict:
    return FONT_CONFIG.get(font_key, FONT_CONFIG["times"])


def get_line_spacing(font_config: dict) -> Pt:
    """Dubbele regelafstand = lettergrootte * 2."""
    return Pt(font_config["size"].pt * 2)


# ---------------------------------------------------------------------------
# Hulpfuncties: XML / opmaak
# ---------------------------------------------------------------------------

def add_field_code(paragraph, field_instruction: str):
    """Voeg een Word-veldcode in (bijv. PAGE of TOC)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_instruction
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)


def add_hyperlink(paragraph, url: str, text: str, font_config: dict):
    """Voeg een klikbare hyperlink toe aan een alinea."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # Stijl: blauw, onderstreept
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    font_elem = OxmlElement("w:rFonts")
    font_elem.set(qn("w:ascii"), font_config["name"])
    font_elem.set(qn("w:hAnsi"), font_config["name"])
    rpr.append(font_elem)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_config["size"].pt * 2)))
    rpr.append(sz)

    run_elem.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def parse_inline_markdown(paragraph, text: str, font_config: dict, extra_bold: bool = False):
    """
    Verwerk **tekst** (vet) en *tekst* (cursief) naar runs in de alinea.
    Ondersteunt ook DOI-URL's als hyperlinks.
    Tekst buiten opmaakmarkeringen is normaal (of vet als extra_bold=True).
    """
    # Splits op ** en * delimiters — ** eerst zodat **bold** niet als *italic* matcht
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            # Vetgedrukte tekst: **tekst**
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.italic = extra_bold
            _apply_font(run, font_config)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            inner = part[1:-1]
            # Controleer op DOI in cursief stuk
            doi_match = re.search(r"(https?://doi\.org/\S+)", inner)
            if doi_match:
                # Tekst voor de DOI cursief
                before = inner[:doi_match.start()]
                doi_url = doi_match.group(1).rstrip(".")
                after = inner[doi_match.end():]
                if before:
                    run = paragraph.add_run(before)
                    run.italic = True
                    run.bold = extra_bold
                    _apply_font(run, font_config)
                add_hyperlink(paragraph, doi_url, doi_url, font_config)
                if after:
                    run = paragraph.add_run(after)
                    run.italic = True
                    run.bold = extra_bold
                    _apply_font(run, font_config)
            else:
                run = paragraph.add_run(inner)
                run.italic = True
                run.bold = extra_bold
                _apply_font(run, font_config)
        else:
            # Controleer op gewone DOI in niet-cursief stuk
            doi_match = re.search(r"(https?://doi\.org/\S+)", part)
            if doi_match:
                before = part[:doi_match.start()]
                doi_url = doi_match.group(1).rstrip(".")
                after = part[doi_match.end():]
                if before:
                    run = paragraph.add_run(before)
                    run.bold = extra_bold
                    _apply_font(run, font_config)
                add_hyperlink(paragraph, doi_url, doi_url, font_config)
                if after:
                    run = paragraph.add_run(after)
                    run.bold = extra_bold
                    _apply_font(run, font_config)
            elif part:
                run = paragraph.add_run(part)
                run.bold = extra_bold
                _apply_font(run, font_config)


def _apply_font(run, font_config: dict):
    """Pas lettertype en -grootte toe op een run."""
    run.font.name = font_config["name"]
    run.font.size = font_config["size"]


def set_paragraph_format(paragraph, font_config: dict, first_line_indent: bool = True,
                          left_indent: Cm = None, hanging_indent: bool = False,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Stel APA-standaard alinea-opmaak in."""
    pf = paragraph.paragraph_format
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    paragraph.alignment = alignment

    if hanging_indent:
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
    elif left_indent is not None:
        pf.left_indent = left_indent
        pf.first_line_indent = Cm(0)
    elif first_line_indent:
        pf.first_line_indent = Cm(1.27)
    else:
        pf.first_line_indent = Cm(0)


def _make_border_elem(name: str, val: str = "single", sz: str = "4") -> OxmlElement:
    """Maak een w:border-element aan."""
    border = OxmlElement(f"w:{name}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")
    return border


def _make_none_border_elem(name: str) -> OxmlElement:
    """Maak een w:border-element aan dat geen lijn tekent."""
    border = OxmlElement(f"w:{name}")
    border.set(qn("w:val"), "none")
    border.set(qn("w:sz"), "0")
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")
    return border


def set_apa_table_borders(table, has_header_row: bool = True):
    """
    Zet APA-conforme tabelranden:
    - Lijn boven de tabel
    - Lijn onder de koptekstrij (alleen als has_header_row=True)
    - Lijn onder de laatste rij
    - Geen andere horizontale of verticale lijnen
    """
    tbl = table._tbl

    # Tabelstijl: verwijder Table Grid overschrijving
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Verwijder bestaande tblBorders als die er al is
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    # Zet tabel-buitenranden: alleen top en bottom
    tblBorders = OxmlElement("w:tblBorders")
    tblBorders.append(_make_border_elem("top"))
    tblBorders.append(_make_none_border_elem("left"))
    tblBorders.append(_make_border_elem("bottom"))
    tblBorders.append(_make_none_border_elem("right"))
    tblBorders.append(_make_none_border_elem("insideH"))
    tblBorders.append(_make_none_border_elem("insideV"))
    tblPr.append(tblBorders)

    # Verwijder celspecifieke randen van alle cellen
    for row_idx, row in enumerate(table.rows):
        is_header = has_header_row and row_idx == 0
        is_last = row_idx == len(table.rows) - 1

        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)

            # Verwijder bestaande tcBorders
            existing_tc = tcPr.find(qn("w:tcBorders"))
            if existing_tc is not None:
                tcPr.remove(existing_tc)

            tcBorders = OxmlElement("w:tcBorders")
            # Alle celranden op none — tabel-niveau regelt top/bottom
            tcBorders.append(_make_none_border_elem("top"))
            tcBorders.append(_make_none_border_elem("left"))
            tcBorders.append(_make_none_border_elem("right"))

            if is_header:
                # Koptekstrij: lijn eronder
                tcBorders.append(_make_border_elem("bottom"))
            elif is_last:
                # Laatste rij: lijn eronder
                tcBorders.append(_make_border_elem("bottom"))
            else:
                tcBorders.append(_make_none_border_elem("bottom"))

            tcPr.append(tcBorders)


def centered_bold_para(doc: Document, text: str, font_config: dict):
    """Voeg een gecentreerde, vetgedrukte alinea toe (Normal-stijl, geen kop)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    _apply_font(run, font_config)
    return p


# ---------------------------------------------------------------------------
# Documentinstellingen
# ---------------------------------------------------------------------------

def _configure_heading_style(doc: Document, style_name: str, font_config: dict,
                              bold: bool = True, italic: bool = False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Pas een Word-kopstijl aan naar APA-opmaak zodat de TOC hem herkent."""
    style = doc.styles[style_name]
    style.font.name = font_config["name"]
    style.font.size = font_config["size"]
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)  # Zwart (geen standaard blauwe kopkleur)
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.page_break_before = False  # Overschrijf Word-standaard True voor H1


def setup_document(doc: Document, font_config: dict):
    """Stel marges, standaard stijl, kopstijlen en paginanummers in."""
    section = doc.sections[0]

    # Marges: 2,54 cm aan alle zijden (APA)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Eerste pagina (titelblad) heeft geen paginanummer
    section.different_first_page_header_footer = True

    # Standaard alineastijl instellen
    style = doc.styles["Normal"]
    style.font.name = font_config["name"]
    style.font.size = font_config["size"]
    pf = style.paragraph_format
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(0)

    # Kopstijlen configureren met APA-opmaak (vereist voor de inhoudsopgave)
    _configure_heading_style(doc, "Heading 1", font_config,
                             bold=True, italic=False,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _configure_heading_style(doc, "Heading 2", font_config,
                             bold=True, italic=False,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _configure_heading_style(doc, "Heading 3", font_config,
                             bold=True, italic=True,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _configure_heading_style(doc, "Heading 4", font_config,
                             bold=True, italic=False,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _configure_heading_style(doc, "Heading 5", font_config,
                             bold=True, italic=True,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # "Code Snippet" stijl aanmaken (Courier New 11pt, grijze achtergrond, 1 cm inspringing)
    try:
        code_style = doc.styles.add_style('Code Snippet', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        # Stijl bestaat al (bijv. bij hergebruik van een template)
        code_style = doc.styles['Code Snippet']
    code_style.base_style = doc.styles['Normal']
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(11)
    code_pf = code_style.paragraph_format
    code_pf.left_indent = Cm(1)
    code_pf.first_line_indent = Cm(0)
    code_pf.space_before = Pt(0)
    code_pf.space_after = Pt(0)
    code_pf.line_spacing = Pt(14)
    code_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Paginanummer rechts in koptekst (vanaf pagina 2)
    header = section.header
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    add_field_code(hp, " PAGE ")

    # Eerste-pagina-koptekst leeg laten (titelblad)
    first_header = section.first_page_header
    if not first_header.paragraphs:
        first_header.add_paragraph()


# ---------------------------------------------------------------------------
# Titelblad
# ---------------------------------------------------------------------------

def build_title_page(doc: Document, metadata: dict, font_config: dict):
    """Bouw het APA-conforme titelblad."""

    def centered_para(text: str, bold: bool = False, font_size: Pt = None, space_before: Pt = Pt(0)) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = space_before
        pf.space_after = Pt(0)
        pf.line_spacing = get_line_spacing(font_config)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = p.add_run(text)
        run.bold = bold
        run.font.name = font_config["name"]
        run.font.size = font_size or font_config["size"]

    # Titelblad: 3 lege regels boven ("drie tot vier regels van boven" — Scribbr NL)
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = get_line_spacing(font_config)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # APA 7 student paper volgorde (Scribbr NL — scribbr.nl/apa-stijl/titelpagina/):
    # 1. Titel (vet) + ondertitel
    centered_para(metadata.get("title", "Zonder Titel"), bold=True)
    if metadata.get("subtitle"):
        centered_para(metadata["subtitle"])

    # 2. Auteurs (lege regel onder titel) + studentnummers
    authors = metadata.get("authors", [])
    student_numbers = metadata.get("student_numbers", [])
    for i, author in enumerate(authors):
        centered_para(author)
        if i < len(student_numbers):
            centered_para(student_numbers[i])

    # 3. Instelling + faculteit
    if metadata.get("institution"):
        centered_para(metadata["institution"])
    if metadata.get("faculty"):
        centered_para(metadata["faculty"])

    # 4. Opleiding, vak, begeleider, datum
    if metadata.get("opleiding"):
        centered_para(metadata["opleiding"])
    if metadata.get("course"):
        centered_para(metadata["course"])
    if metadata.get("supervisor"):
        centered_para(metadata["supervisor"])
    if metadata.get("submission_date"):
        centered_para(metadata["submission_date"])

    # Geen doc.add_page_break() — volgende sectie start via page_break_before op H1


# ---------------------------------------------------------------------------
# Samenvatting
# ---------------------------------------------------------------------------

def build_abstract(doc: Document, abstract: dict, font_config: dict):
    """Bouw de samenvattingspagina (APA: 'Samenvatting')."""
    build_section_heading(doc, "Samenvatting", font_config)

    # Abstracttekst: geen eerste-regel-inspringing (APA)
    p = doc.add_paragraph()
    set_paragraph_format(p, font_config, first_line_indent=False)
    run = p.add_run(abstract.get("text", ""))
    _apply_font(run, font_config)

    # Sleutelwoorden (APA: géén eerste-regels-inspringing bij keywords-label)
    keywords = abstract.get("keywords", [])
    if keywords:
        p = doc.add_paragraph()
        set_paragraph_format(p, font_config, first_line_indent=False)
        kw_run = p.add_run("Sleutelwoorden: ")
        kw_run.italic = True
        _apply_font(kw_run, font_config)
        rest_run = p.add_run(", ".join(keywords))
        _apply_font(rest_run, font_config)


# ---------------------------------------------------------------------------
# Inhoudsopgave
# ---------------------------------------------------------------------------

def build_toc(doc: Document, font_config: dict):
    """
    Voeg een inhoudsopgave in met een Word TOC-veld.
    De gebruiker kan de inhoudsopgave bijwerken via:
    rechtermuisknop op de inhoudsopgave → 'Veld bijwerken'.
    """
    build_section_heading(doc, "Inhoudsopgave", font_config)

    # TOC-veld alinea
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_before = Pt(0)
    toc_para.paragraph_format.space_after = Pt(0)

    # Bouw het TOC-veld handmatig via XML
    p_elem = toc_para._p
    run_elem = OxmlElement("w:r")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run_elem.append(fld_begin)
    p_elem.append(run_elem)

    run_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr.append(instr)
    p_elem.append(run_instr)

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)
    p_elem.append(run_end)


# ---------------------------------------------------------------------------
# Lijst van Afkortingen
# ---------------------------------------------------------------------------

def build_abbreviations(doc: Document, abbreviations: list, font_config: dict):
    """Bouw de lijst van afkortingen (alleen als er afkortingen zijn)."""
    if not abbreviations:
        return

    build_section_heading(doc, "Lijst van Afkortingen", font_config)

    # Tabel (geen koptekstrij; APA: alleen top+bottom border)
    table = doc.add_table(rows=len(abbreviations), cols=2)

    for i, item in enumerate(abbreviations):
        row = table.rows[i]
        # Kolom 1: afkorting (vet)
        cell_a = row.cells[0]
        cell_a.text = ""
        p = cell_a.paragraphs[0]
        set_paragraph_format(p, font_config, first_line_indent=False)
        run = p.add_run(item.get("afkorting", ""))
        run.bold = True
        _apply_font(run, font_config)

        # Kolom 2: definitie
        cell_b = row.cells[1]
        cell_b.text = ""
        p = cell_b.paragraphs[0]
        set_paragraph_format(p, font_config, first_line_indent=False)
        run = p.add_run(item.get("definitie", ""))
        _apply_font(run, font_config)

    set_apa_table_borders(table, has_header_row=False)


# ---------------------------------------------------------------------------
# Blokrenderers (gedeeld door inleiding, body, conclusie, bijlagen)
# ---------------------------------------------------------------------------

def render_paragraph(doc: Document, block: dict, font_config: dict):
    """Render een normale alinea. Ondersteunt \\n als zachte regelafbreking binnen de alinea."""
    p = doc.add_paragraph()
    set_paragraph_format(p, font_config, first_line_indent=True)
    text = block.get("text", "")
    if "\n" in text:
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()  # zachte regelafbreking (w:br)
            parse_inline_markdown(p, line, font_config)
    else:
        parse_inline_markdown(p, text, font_config)


def render_heading(doc: Document, block: dict, font_config: dict):
    """Render een kop (niveau 1–5) met Word-kopstijl zodat de TOC werkt."""
    level = block.get("level", 1)
    text = block.get("text", "")
    inline_text = block.get("inline_text", "")

    if level in (1, 2, 3):
        # Gebruik Word-kopstijl → TOC herkent deze automatisch
        p = doc.add_paragraph(style=f"Heading {level}")
        p.paragraph_format.keep_with_next = True
        if level == 1:
            # H1 in de body start altijd op een nieuwe pagina
            p.paragraph_format.page_break_before = True
        run = p.add_run(text)
        _apply_font(run, font_config)

    elif level == 4:
        # Koptekst vet + punt inline met alineatekst; gebruik Heading 4-stijl
        p = doc.add_paragraph(style="Heading 4")
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text + ". ")
        _apply_font(run, font_config)
        if inline_text:
            # Inline tekst heeft normale opmaak (niet vet)
            inline_run = p.add_run(inline_text)
            inline_run.bold = False
            inline_run.italic = False
            _apply_font(inline_run, font_config)

    elif level == 5:
        # Koptekst vet cursief + punt inline met alineatekst; gebruik Heading 5-stijl
        p = doc.add_paragraph(style="Heading 5")
        p.paragraph_format.first_line_indent = Cm(1.27)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text + ". ")
        _apply_font(run, font_config)
        if inline_text:
            inline_run = p.add_run(inline_text)
            inline_run.bold = False
            inline_run.italic = False
            _apply_font(inline_run, font_config)


def _set_paragraph_shading(paragraph, fill_color: str):
    """Zet achtergrondkleur op een alinea via pPr/shd XML."""
    pPr = paragraph._p.get_or_add_pPr()
    # Verwijder bestaande shd-elementen
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    pPr.append(shd)


def render_code(doc: Document, block: dict, font_config: dict):
    """
    Render een codeblok of terminal-output.
    Opmaak: Courier New 11pt, grijze achtergrond (#D9D9D9), 1 cm links inspringen, enkelvoudig.
    Elke regel = aparte alinea zodat de achtergrond ononderbroken doorloopt.
    """
    tekst = block.get("tekst", block.get("text", ""))
    lines = tekst.split("\n") if tekst else [""]

    for line in lines:
        if not line:
            continue  # lege regels in codeblok overslaan — voorkomt grijze gaten
        p = doc.add_paragraph(style='Code Snippet')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_shading(p, "D9D9D9")

        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(11)


def render_block_quote(doc: Document, block: dict, font_config: dict):
    """Render een blokcitaat (≥40 woorden, APA)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(0)
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(block.get("text", ""))
    _apply_font(run, font_config)

    citation = block.get("citation", "")
    if citation:
        run2 = p.add_run(" " + citation)
        _apply_font(run2, font_config)


def render_list_item(doc: Document, block: dict, font_config: dict):
    """Render een lijstitem met bullet-teken en inspringing."""
    p = doc.add_paragraph()
    text = block.get("text", "")
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-0.63)
    pf.line_spacing = get_line_spacing(font_config)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("\u2022 ")
    _apply_font(run, font_config)
    parse_inline_markdown(p, text, font_config)


def render_table(doc: Document, block: dict, font_config: dict):
    """Render een APA-tabel: label + titel bóven, geen verticale lijnen, noot ónder."""
    number = block.get("number")
    title = block.get("title", "")
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    note = block.get("note", "")

    # Tabelnummer + titel alleen als expliciet opgegeven (voorkomt "Tabel 1" op niet-APA tabellen)
    if number is not None:
        p_num = doc.add_paragraph()
        set_paragraph_format(p_num, font_config, first_line_indent=False, left_indent=Cm(0))
        run = p_num.add_run(f"Tabel {number}")
        run.bold = True
        _apply_font(run, font_config)

    if title:
        p_title = doc.add_paragraph()
        set_paragraph_format(p_title, font_config, first_line_indent=False, left_indent=Cm(0))
        run = p_title.add_run(title)
        run.italic = True
        _apply_font(run, font_config)

    # Validatie: headers verplicht, rijen moeten kloppen met kolomaantal
    tbl_label = f"Tabel {number}" if number is not None else "Tabel"
    if not headers:
        msg = f"WAARSCHUWING: {tbl_label} heeft geen headers — tabel overgeslagen."
        _build_warnings.append(msg)
        print(msg, file=sys.stderr)
        return
    n_cols = len(headers)
    validated_rows = []
    for i, row_data in enumerate(rows):
        if len(row_data) != n_cols:
            msg = (f"WAARSCHUWING: {tbl_label} rij {i + 1} heeft {len(row_data)} cel(len), "
                   f"verwacht {n_cols}. Rij wordt aangevuld of afgekapt.")
            _build_warnings.append(msg)
            print(msg, file=sys.stderr)
            # Aanvullen met lege cellen of afkappen
            row_data = list(row_data)[:n_cols] + [""] * max(0, n_cols - len(row_data))
        validated_rows.append(row_data)
    rows = validated_rows

    # Tabel
    total_rows = len(rows) + 1  # +1 voor koptekstrij
    table = doc.add_table(rows=total_rows, cols=n_cols)
    # Geen tabelstijl instellen (standaard = geen borders; APA-borders via set_apa_table_borders)
    # Koptekstrij
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, font_config, first_line_indent=False)
        run = p.add_run(hdr)
        run.bold = True
        _apply_font(run, font_config)

    # Gegevensrijen
    for i, row_data in enumerate(rows):
        data_row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            cell = data_row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, font_config, first_line_indent=False)
            parse_inline_markdown(p, str(cell_text), font_config)

    set_apa_table_borders(table, has_header_row=True)

    # Tabelvoetnoot — "Noot." cursief + gewone tekst
    if note:
        p_note = doc.add_paragraph()
        set_paragraph_format(p_note, font_config, first_line_indent=False)
        run_label = p_note.add_run("Noot.")
        run_label.italic = True
        _apply_font(run_label, font_config)
        run_text = p_note.add_run(" " + note)
        _apply_font(run_text, font_config)


def render_figure_placeholder(doc: Document, block: dict, font_config: dict):
    """Render een figuur met label + caption bóven de afbeelding (APA 7)."""
    number = block.get("number", 1)
    caption = block.get("caption", "")
    note = block.get("note", "")
    image_path = block.get("image_path", "")

    # APA 7: Figuur N. (vet) + caption (cursief) BOVEN de afbeelding
    p_label = doc.add_paragraph()
    set_paragraph_format(p_label, font_config, first_line_indent=False, left_indent=Cm(0))
    run_num = p_label.add_run(f"Figuur {number}")
    run_num.bold = True
    _apply_font(run_num, font_config)
    if caption:
        p_caption = doc.add_paragraph()
        set_paragraph_format(p_caption, font_config, first_line_indent=False, left_indent=Cm(0))
        run_cap = p_caption.add_run(caption)
        run_cap.italic = True
        _apply_font(run_cap, font_config)

    # Afbeelding of plaatshouder
    if image_path:
        try:
            p_img = doc.add_paragraph()
            set_paragraph_format(p_img, font_config, first_line_indent=False,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
            p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            p_img.paragraph_format.line_spacing = Pt(2)
            run_img = p_img.add_run()
            # Bewaar originele beeldverhouding: schaal alleen als breder dan 5.5 inch
            img_width = Inches(5.5)
            try:
                from PIL import Image as PILImage
                with PILImage.open(image_path) as pil_img:
                    w_px, h_px = pil_img.size
                    dpi = pil_img.info.get('dpi', (96, 96))
                    w_inch = w_px / dpi[0]
                    if w_inch < 5.5:
                        img_width = Inches(w_inch)
            except Exception:
                pass  # fallback naar 5.5 inch
            run_img.add_picture(image_path, width=img_width)
        except FileNotFoundError:
            msg = f"WAARSCHUWING: afbeelding niet gevonden — {image_path}. Plaatshouder ingevoegd."
            print(msg, file=sys.stderr)
            _build_warnings.append(msg)
            p_placeholder = doc.add_paragraph()
            set_paragraph_format(p_placeholder, font_config, first_line_indent=False,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p_placeholder.add_run(f"[Figuur hier invoegen: {image_path}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            _apply_font(run, font_config)
    else:
        p_placeholder = doc.add_paragraph()
        set_paragraph_format(p_placeholder, font_config, first_line_indent=False,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p_placeholder.add_run("[Figuur hier invoegen]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        _apply_font(run, font_config)

    # Figuurvoetnoot (onder de afbeelding)
    if note:
        p_note = doc.add_paragraph()
        set_paragraph_format(p_note, font_config, first_line_indent=False)
        run_label = p_note.add_run("Noot.")
        run_label.italic = True
        _apply_font(run_label, font_config)
        run_text = p_note.add_run(" " + note)
        _apply_font(run_text, font_config)


def render_block(doc: Document, block: dict, font_config: dict):
    """Dispatcher: verwerk één blok op basis van het type."""
    block_type = block.get("type", "paragraph")

    if block_type == "paragraph":
        render_paragraph(doc, block, font_config)
    elif block_type == "heading":
        render_heading(doc, block, font_config)
    elif block_type == "block_quote":
        render_block_quote(doc, block, font_config)
    elif block_type == "table":
        render_table(doc, block, font_config)
    elif block_type == "figure_placeholder":
        render_figure_placeholder(doc, block, font_config)
    elif block_type == "code":
        render_code(doc, block, font_config)
    elif block_type == "list_item":
        render_list_item(doc, block, font_config)
    elif block_type == "page_break":
        doc.add_page_break()
    else:
        # Onbekend type: sla over maar log naar stderr
        print(f"Waarschuwing: onbekend bloktype '{block_type}' overgeslagen.", file=sys.stderr)


# ---------------------------------------------------------------------------
# Sectiebouwers
# ---------------------------------------------------------------------------

def build_section_heading(doc: Document, text: str, font_config: dict):
    """Kop niveau 1 (gecentreerd, vet) als sectie-opener — met Word Heading 1-stijl voor de TOC."""
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    _apply_font(run, font_config)


def build_introduction(doc: Document, title: str, intro_blocks: list, font_config: dict):
    """Bouw de inleiding: kop = documenttitel (APA-regel)."""
    build_section_heading(doc, title, font_config)
    for block in intro_blocks:
        render_block(doc, block, font_config)


def build_body(doc: Document, body_blocks: list, font_config: dict):
    """Render alle body-blokken."""
    for block in body_blocks:
        render_block(doc, block, font_config)


def build_conclusion(doc: Document, conclusion_blocks: list, font_config: dict):
    """Bouw de conclusie: kop 'Conclusie' (altijd aanwezig)."""
    build_section_heading(doc, "Conclusie", font_config)
    for block in conclusion_blocks:
        render_block(doc, block, font_config)


def build_reference_list(doc: Document, references: list, font_config: dict):
    """
    Bouw de literatuurlijst: hangende inspringing, cursief via *...*, DOI als hyperlink.
    Kop: 'Literatuurlijst' (kop niveau 1, apart pagina).
    """
    build_section_heading(doc, "Literatuurlijst", font_config)

    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, font_config, hanging_indent=True)
        parse_inline_markdown(p, ref, font_config)


def build_appendices(doc: Document, appendices: list, font_config: dict):
    """
    Bouw de bijlagensectie conform APA/Scribbr.
    Geen overkoepelende 'Bijlagen'-kop. Elke bijlage start op een nieuwe pagina.
    'Bijlage A' is Heading 1 (verschijnt in TOC), titel eronder als gecentreerde normale alinea.
    """
    if not appendices:
        return

    for i, appendix in enumerate(appendices):
        # Bijlage-label (bijv. "Bijlage A"): Heading 1 → verschijnt in TOC
        label = appendix.get("label", f"Bijlage {chr(65 + i)}")
        build_section_heading(doc, label, font_config)

        # Beschrijvende titel: gecentreerd vet, Normal-stijl (niet in TOC)
        if appendix.get("title"):
            centered_bold_para(doc, appendix["title"], font_config)

        # Inhoud
        for block in appendix.get("content", []):
            render_block(doc, block, font_config)


# ---------------------------------------------------------------------------
# Hoofdorchestrator
# ---------------------------------------------------------------------------

def build_document(data: dict, output_path: str) -> str:
    """Bouw het volledige APA-document en sla op. Geeft het outputpad terug."""
    metadata = data.get("metadata", {})
    font_key = metadata.get("font", "times")
    font_config = get_font_config(font_key)

    doc = Document()

    # 1. Documentinstelling (marges, stijlen, paginanummers)
    setup_document(doc, font_config)

    # 2. Titelblad
    build_title_page(doc, metadata, font_config)

    # 3. Samenvatting
    abstract = data.get("abstract", {})
    if abstract:
        build_abstract(doc, abstract, font_config)

    # 4. Inhoudsopgave (altijd aanwezig)
    build_toc(doc, font_config)

    # 5. Lijst van Afkortingen (alleen als er afkortingen zijn)
    abbreviations = data.get("abbreviations", [])
    if abbreviations:
        build_abbreviations(doc, abbreviations, font_config)

    # 6. Inleiding (kop = documenttitel, APA) — alleen als er inhoud is
    intro_blocks = data.get("introduction_text", [])
    if intro_blocks:
        build_introduction(doc, metadata.get("title", "Inleiding"), intro_blocks, font_config)

    # 7. Body secties
    body_blocks = data.get("body", [])
    if body_blocks:
        build_body(doc, body_blocks, font_config)

    # 8. Conclusie — alleen als er inhoud is
    conclusion_blocks = data.get("conclusion_text", [])
    if conclusion_blocks:
        build_conclusion(doc, conclusion_blocks, font_config)

    # 9. Literatuurlijst (altijd aanwezig)
    references = data.get("references", [])
    build_reference_list(doc, references, font_config)

    # 10. Bijlagen (altijd aanwezig)
    appendices = data.get("appendices", [])
    build_appendices(doc, appendices, font_config)

    # Opslaan
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


# ---------------------------------------------------------------------------
# Validatie
# ---------------------------------------------------------------------------

def validate_data(data: dict):
    """Controleer verplichte velden in de JSON-payload."""
    errors = []
    if "metadata" not in data:
        errors.append("Veld 'metadata' ontbreekt.")
    else:
        meta = data["metadata"]
        if not meta.get("title"):
            errors.append("metadata.title is verplicht.")

    font = data.get("metadata", {}).get("font", "times")
    if font not in FONT_CONFIG:
        errors.append(f"Ongeldig lettertype '{font}'. Kies uit: times, arial, georgia.")

    if errors:
        raise ValueError("Validatiefouten:\n" + "\n".join(f"  - {e}" for e in errors))


# ---------------------------------------------------------------------------
# CLI-ingang
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Genereer een APA 7-conform Word-document (.docx) op basis van een JSON-payload."
    )
    parser.add_argument("--input", required=True, help="Pad naar de JSON-invoer (.tmp/export_payload.json)")
    parser.add_argument("--output", required=False, help="Pad voor het uitvoer-.docx-bestand")
    args = parser.parse_args()

    # Laad JSON
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Fout: invoerbestand '{args.input}' niet gevonden.", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Fout: ongeldige JSON in '{args.input}': {e}", file=sys.stderr)
        sys.exit(1)

    # Valideer
    try:
        validate_data(data)
    except ValueError as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)

    # Bepaal outputpad
    if args.output:
        output_path = args.output
    else:
        title = data.get("metadata", {}).get("title", "rapport")
        slug = re.sub(r"[^\w\s-]", "", title.lower())
        slug = re.sub(r"[\s-]+", "_", slug)[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = os.path.join(".tmp", f"rapport_{slug}_{timestamp}.docx")

    # Genereer document
    _build_warnings.clear()
    try:
        result_path = build_document(data, output_path)
        for w in _build_warnings:
            print(w)
        print(result_path)
    except Exception as e:
        print(f"Fout bij het genereren van het document: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
